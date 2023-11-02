from math import e
import threading

from docx import Document

from parsers.parser import Parser
from runner.run import Runner
from translator.translator import Translator


class DocxParser(Parser):
    def __init__(self, in_path, out_path) -> None:
        self.in_document = Document(in_path)
        self.translator = Translator()
        self.runner = Runner()
        self.lock = threading.Lock()
        self.out_path = out_path

    def run(self) -> None:
        print("Processing paragraphs...")

        self.runner.run(self.process_element, self.in_document.paragraphs)
        print("Processing tables...")
        tables = []
        for table in self.in_document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        tables.append(paragraph)
        try:
            self.runner.run(self.process_element, tables)
        except Exception as e:
            print(e)

        print(f"Saving at {self.out_path}")
        self.in_document.save(self.out_path)

    def process_element(self, element):
        cn_text = element.text
        en_text = self.translator.cn2en(cn_text)

        if len(element.runs) > 0:
            with self.lock:
                element.runs[0].text = en_text
                for i, r in enumerate(element.runs):
                    if i != 0:
                        r.text = ""
